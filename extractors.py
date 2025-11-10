"""
Content extractors for PDF and DOC files.
All extractors return multimodal content (text, images, tables, etc.).

Supported file types:
    - PDF (.pdf)
    - Microsoft Word (.docx)

Extractors focus solely on content extraction from files.
Processing stages (OCR, chunking, etc.) are in the stages/ directory.
"""

import os
import tempfile
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
import dtlpy as dl

# ============================================================================
# MULTIMODAL DATA STRUCTURES
# ============================================================================


@dataclass
class ImageContent:
    """Represents an extracted image"""

    path: str  # File path or temp path
    caption: Optional[str] = None  # Image caption/alt text
    page_number: Optional[int] = None  # Source page (for PDFs)
    bbox: Optional[tuple] = None  # Bounding box (x, y, w, h)
    format: Optional[str] = None  # png, jpg, etc.
    size: Optional[tuple] = None  # (width, height)
    data: Optional[bytes] = None  # Raw image bytes (optional)

    def to_dict(self):
        return {
            'path': self.path,
            'caption': self.caption,
            'page_number': self.page_number,
            'bbox': self.bbox,
            'format': self.format,
            'size': self.size,
        }


@dataclass
class TableContent:
    """Represents an extracted table"""

    data: Any  # pandas DataFrame or list of dicts
    markdown: Optional[str] = None  # Table as markdown
    html: Optional[str] = None  # Table as HTML
    page_number: Optional[int] = None  # Source page
    location: Optional[Dict] = None  # Position info

    def to_dict(self):
        return {
            'markdown': self.markdown,
            'html': self.html,
            'page_number': self.page_number,
            'location': self.location,
        }


@dataclass
class AudioContent:
    """Represents audio content"""

    path: str  # File path
    duration: Optional[float] = None  # Duration in seconds
    transcript: Optional[str] = None  # Transcription
    format: Optional[str] = None  # mp3, wav, etc.
    sample_rate: Optional[int] = None  # Sample rate

    def to_dict(self):
        return {
            'path': self.path,
            'duration': self.duration,
            'transcript': self.transcript,
            'format': self.format,
            'sample_rate': self.sample_rate,
        }


@dataclass
class ExtractedContent:
    """
    Multimodal content extracted from document.
    Can contain text, images, tables, audio, etc.
    """

    text: str = ""  # Main text content
    images: List[ImageContent] = field(default_factory=list)
    tables: List[TableContent] = field(default_factory=list)
    audio: List[AudioContent] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self):
        """Convert to dictionary for pipeline processing"""
        return {
            'content': self.text,
            'images': [img.to_dict() for img in self.images],
            'tables': [tbl.to_dict() for tbl in self.tables],
            'audio': [aud.to_dict() for aud in self.audio],
            'metadata': self.metadata,
        }

    def has_images(self) -> bool:
        return len(self.images) > 0

    def has_tables(self) -> bool:
        return len(self.tables) > 0

    def has_audio(self) -> bool:
        return len(self.audio) > 0


# ============================================================================
# BASE EXTRACTOR
# ============================================================================


class BaseExtractor(ABC):
    """Base class for all content extractors"""

    def __init__(self, mime_type: str, name: str):
        self.mime_type = mime_type
        self.name = name

    @abstractmethod
    def extract(self, item: dl.Item, config: Dict[str, Any]) -> ExtractedContent:
        """
        Extract multimodal content from Dataloop item.

        Args:
            item: Dataloop item to process
            config: Configuration dictionary

        Returns:
            ExtractedContent: Object containing text, images, tables, etc.
        """
        pass

    def __repr__(self):
        return f"{self.name}Extractor('{self.mime_type}')"


# ============================================================================
# PDF EXTRACTOR
# ============================================================================


class PDFExtractor(BaseExtractor):
    """Extract text, images, and tables from PDF files"""

    def __init__(self):
        super().__init__('application/pdf', 'PDF')

    def extract(self, item: dl.Item, config: Dict[str, Any]) -> ExtractedContent:
        """Extract all content from PDF"""
        import fitz  # PyMuPDF

        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = item.download(local_path=temp_dir)

            # Choose extraction method
            if config.get('use_markdown_extraction', False):
                return self._extract_with_markdown(file_path, item, temp_dir, config)
            else:
                return self._extract_with_pymupdf(file_path, item, temp_dir, config)

    def _extract_with_pymupdf(
        self, file_path: str, item: dl.Item, temp_dir: str, config: Dict[str, Any]
    ) -> ExtractedContent:
        """Extract using basic PyMuPDF"""
        import fitz

        doc = fitz.open(file_path)
        result = ExtractedContent()
        text_parts = []

        for page_num, page in enumerate(doc):
            # Extract text
            page_text = page.get_text()
            text_parts.append(f"\n\n--- Page {page_num + 1} ---\n\n{page_text}")

            # Extract images if requested
            if config.get('extract_images', True):
                images = self._extract_images_from_page(page, page_num, temp_dir)
                result.images.extend(images)

        result.text = ''.join(text_parts)
        result.metadata = {
            'page_count': len(doc),
            'source_file': item.name,
            'extraction_method': 'pymupdf',
            'image_count': len(result.images),
            'table_count': len(result.tables),
            'extractor': 'pdf',
        }

        doc.close()
        return result

    def _extract_with_markdown(
        self, file_path: str, item: dl.Item, temp_dir: str, config: Dict[str, Any]
    ) -> ExtractedContent:
        """Extract using pymupdf4llm (preserves structure)"""
        import pymupdf4llm
        import fitz

        # Extract as markdown
        md_text = pymupdf4llm.to_markdown(file_path)

        result = ExtractedContent()
        result.text = md_text

        # Still extract images if requested
        if config.get('extract_images', True):
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                images = self._extract_images_from_page(page, page_num, temp_dir)
                result.images.extend(images)
            doc.close()

        result.metadata = {
            'source_file': item.name,
            'extraction_method': 'pymupdf4llm',
            'format': 'markdown',
            'image_count': len(result.images),
            'extractor': 'pdf',
        }

        return result

    def _extract_images_from_page(self, page, page_num: int, temp_dir: str) -> List[ImageContent]:
        """
        Extract images from a PDF page with positional metadata.

        Extracts images along with their bounding box positions on the page.
        """
        images = []
        image_list = page.get_images(full=True)

        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                base_image = page.parent.extract_image(xref)

                image_path = f"{temp_dir}/page{page_num}_img{img_index}.{base_image['ext']}"
                with open(image_path, 'wb') as f:
                    f.write(base_image['image'])

                # Get image bounding boxes/positions on the page
                bbox = None
                image_rects = page.get_image_rects(xref)
                if image_rects:
                    # Use the first (or largest) rectangle if multiple found
                    rect = image_rects[0] if isinstance(image_rects, list) else image_rects
                    # Convert fitz.Rect to (x0, y0, x1, y1) then to (x, y, width, height)
                    bbox = (rect.x0, rect.y0, rect.width, rect.height)  # x position  # y position  # width  # height

                images.append(
                    ImageContent(
                        path=image_path,
                        page_number=page_num + 1,
                        format=base_image['ext'],
                        size=(base_image.get('width'), base_image.get('height')),
                        bbox=bbox,
                    )
                )
            except Exception as e:
                print(f"Warning: Failed to extract image {img_index} from page {page_num}: {e}")

        return images


# ============================================================================
# DOCS EXTRACTOR (Google Docs as .docx)
# ============================================================================


class DocsExtractor(BaseExtractor):
    """Extract text and images from Google Docs (.docx)"""

    def __init__(self):
        super().__init__('application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'Docs')

    def extract(self, item: dl.Item, config: Dict[str, Any]) -> ExtractedContent:
        """Extract content from .docx"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = item.download(local_path=temp_dir)
            doc = Document(file_path)

            result = ExtractedContent()

            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            result.text = '\n\n'.join(paragraphs)

            # Extract images if requested
            if config.get('extract_images', True):
                result.images = self._extract_images(doc, temp_dir)

            # Extract tables if requested
            if config.get('extract_tables', False):
                result.tables = self._extract_tables(doc)

            result.metadata = {
                'paragraph_count': len(paragraphs),
                'source_file': item.name,
                'image_count': len(result.images),
                'table_count': len(result.tables),
                'extractor': 'docs',
            }

            return result

    def _extract_images(self, doc, temp_dir: str) -> List[ImageContent]:
        """Extract embedded images from .docx"""
        images = []

        try:
            # .docx images are in doc.part.rels
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_path = f"{temp_dir}/{rel.target_ref.split('/')[-1]}"
                    with open(image_path, 'wb') as f:
                        f.write(rel.target_part.blob)

                    images.append(
                        ImageContent(
                            path=image_path, format=rel.target_ref.split('.')[-1] if '.' in rel.target_ref else None
                        )
                    )
        except Exception as e:
            print(f"Warning: Failed to extract images from .docx: {e}")

        return images

    def _extract_tables(self, doc) -> List[TableContent]:
        """Extract tables from .docx"""
        tables = []

        for table in doc.tables:
            try:
                # Convert to list of dicts
                rows = []
                headers = [cell.text for cell in table.rows[0].cells]

                for row in table.rows[1:]:
                    row_data = {headers[i]: cell.text for i, cell in enumerate(row.cells)}
                    rows.append(row_data)

                # Convert to markdown
                markdown = self._table_to_markdown(headers, rows)

                tables.append(TableContent(data=rows, markdown=markdown))
            except Exception as e:
                print(f"Warning: Failed to extract table: {e}")

        return tables

    def _table_to_markdown(self, headers: List[str], rows: List[Dict]) -> str:
        """Convert table to markdown format"""
        md = "| " + " | ".join(headers) + " |\n"
        md += "| " + " | ".join(["---"] * len(headers)) + " |\n"

        for row in rows:
            md += "| " + " | ".join([str(row.get(h, '')) for h in headers]) + " |\n"

        return md


# ============================================================================
# REGISTRY
# ============================================================================

EXTRACTOR_REGISTRY = {
    # PDF
    'application/pdf': PDFExtractor,
    # Google Docs / Word
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocsExtractor,
    'application/vnd.google-apps.document': DocsExtractor,
}


def get_extractor(mime_type: str) -> BaseExtractor:
    """
    Get extractor for MIME type.

    Args:
        mime_type: MIME type string

    Returns:
        Extractor instance

    Raises:
        ValueError: If MIME type not supported
    """
    extractor_class = EXTRACTOR_REGISTRY.get(mime_type)

    if not extractor_class:
        raise ValueError(f"Unsupported MIME type: {mime_type}\n" f"Supported types: {list(EXTRACTOR_REGISTRY.keys())}")

    return extractor_class()


def register_extractor(mime_type: str, extractor_class):
    """Register custom extractor"""
    EXTRACTOR_REGISTRY[mime_type] = extractor_class


def get_supported_types():
    """Get list of supported MIME types"""
    return list(EXTRACTOR_REGISTRY.keys())
