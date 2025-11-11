"""
DOCX extractor for extracting text, images, and tables from Word documents.
"""

import tempfile
import os
from typing import List, Dict, Any
import dtlpy as dl
from docx import Document
from .content_types import ExtractedContent, ImageContent, TableContent


class DocsExtractor:
    """Extract text and images from Google Docs (.docx)"""

    def __init__(self):
        self.mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        self.name = 'Docs'

    def extract(self, item: dl.Item, config: Dict[str, Any]) -> ExtractedContent:
        """Extract content from .docx"""
        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = item.download(local_path=temp_dir)
            doc = Document(file_path)

            result = ExtractedContent()

            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            result.text = '\n\n'.join(paragraphs)

            # Extract images if requested
            if config.get('extract_images', True):
                result.images = self._extract_images(doc, temp_dir)

            # Extract tables if requested
            if config.get('extract_tables', False):
                result.tables = self._extract_tables(doc)

            result.metadata = {
                'paragraph_count': len(paragraphs),
                'source_file': item.name,
                'image_count': len(result.images),
                'table_count': len(result.tables),
                'extractor': 'docs',
            }

            return result

    def _extract_images(self, doc, temp_dir: str) -> List[ImageContent]:
        """Extract embedded images from .docx"""
        images = []

        try:
            # .docx images are in doc.part.rels
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_path = os.path.join(temp_dir, rel.target_ref.split('/')[-1])
                    with open(image_path, 'wb') as f:
                        f.write(rel.target_part.blob)

                    images.append(
                        ImageContent(
                            path=image_path, format=rel.target_ref.split('.')[-1] if '.' in rel.target_ref else None
                        )
                    )
        except Exception as e:
            print(f"Warning: Failed to extract images from .docx: {e}")

        return images

    def _extract_tables(self, doc) -> List[TableContent]:
        """Extract tables from .docx"""
        tables = []

        for table in doc.tables:
            try:
                # Convert to list of dicts
                rows = []
                headers = [cell.text for cell in table.rows[0].cells]

                for row in table.rows[1:]:
                    row_data = {headers[i]: cell.text for i, cell in enumerate(row.cells)}
                    rows.append(row_data)

                # Convert to markdown
                markdown = self._table_to_markdown(headers, rows)

                tables.append(TableContent(data=rows, markdown=markdown))
            except Exception as e:
                print(f"Warning: Failed to extract table: {e}")

        return tables

    def _table_to_markdown(self, headers: List[str], rows: List[Dict]) -> str:
        """Convert table to markdown format"""
        md = "| " + " | ".join(headers) + " |\n"
        md += "| " + " | ".join(["---"] * len(headers)) + " |\n"

        for row in rows:
            md += "| " + " | ".join([str(row.get(h, '')) for h in headers]) + " |\n"

        return md
