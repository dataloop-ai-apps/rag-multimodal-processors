"""
DOCX extraction logic.

Handles DOCX-specific extraction operations:
- Text extraction from paragraphs
- Image extraction from embedded resources
- Table extraction with markdown conversion
"""

import io
import logging
import os
import tempfile
from typing import List, Dict, Tuple, Optional

from docx import Document
from PIL import Image

from utils.extracted_data import ExtractedData
from utils.data_types import ImageContent, TableContent

logger = logging.getLogger("rag-preprocessor")


class DOCExtractor:
    """DOCX extraction operations."""

    @staticmethod
    def extract(data: ExtractedData) -> ExtractedData:
        """Extract content from DOCX item."""
        data.current_stage = "extraction"

        if not data.item:
            data.log_error("No item provided for extraction")
            return data

        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                file_path = data.item.download(local_path=temp_dir)
                doc = Document(file_path)

                # Extract images if configured
                if data.config.extract_images:
                    data.images = DOCExtractor._extract_images(doc, temp_dir)

                # Extract tables if configured
                if data.config.extract_tables:
                    data.tables = DOCExtractor._extract_tables(doc)

                # Extract content as markdown
                data.content_text = DOCExtractor._extract_markdown(doc, data.tables)

                # Set metadata
                data.metadata = {
                    'source_file': data.item_name,
                    'extraction_method': 'python-docx',
                    'format': 'markdown',
                    'image_count': len(data.images),
                    'table_count': len(data.tables),
                    'processor': 'doc',
                }

        except Exception as e:
            data.log_error("Document extraction failed. Check logs for details.")
            logger.exception(f"DOCX extraction error: {e}")

        return data

    @staticmethod
    def _extract_markdown(doc: Document, tables: List[TableContent]) -> str:
        """Convert DOCX content to markdown format."""
        md_parts = []
        table_iter = iter(tables)
        current_table = next(table_iter, None)

        for element in doc.element.body:
            tag = element.tag.split('}')[-1]

            if tag == 'p':
                para = doc.paragraphs[DOCExtractor._get_para_index(doc, element)]
                md_line = DOCExtractor._paragraph_to_markdown(para)
                if md_line:
                    md_parts.append(md_line)

            elif tag == 'tbl' and current_table:
                # Insert table markdown inline
                md_parts.append(current_table.markdown)
                current_table = next(table_iter, None)

        return '\n\n'.join(md_parts)

    @staticmethod
    def _get_para_index(doc: Document, element) -> int:
        """Get paragraph index for an element."""
        for i, para in enumerate(doc.paragraphs):
            if para._element is element:
                return i
        return 0

    @staticmethod
    def _paragraph_to_markdown(para) -> str:
        """Convert a paragraph to markdown."""
        text = para.text.strip()
        if not text:
            return ""

        style_name = para.style.name.lower() if para.style else ""

        # Handle headings
        if 'heading' in style_name:
            level = DOCExtractor._get_heading_level(style_name)
            return f"{'#' * level} {text}"

        # Handle list items
        if DOCExtractor._is_list_item(para):
            if DOCExtractor._is_numbered_list(para):
                return f"1. {text}"
            return f"- {text}"

        # Handle regular paragraph with inline formatting
        return DOCExtractor._apply_inline_formatting(para)

    @staticmethod
    def _get_heading_level(style_name: str) -> int:
        """Extract heading level from style name."""
        for i in range(1, 10):
            if str(i) in style_name:
                return i
        return 1

    @staticmethod
    def _is_list_item(para) -> bool:
        """Check if paragraph is a list item."""
        pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            numPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            return numPr is not None
        return False

    @staticmethod
    def _is_numbered_list(para) -> bool:
        """Check if list item is numbered."""
        pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            numPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is not None:
                numId = numPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
                if numId is not None:
                    val = numId.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    # numId > 1 typically indicates numbered list
                    return val and int(val) > 1
        return False

    @staticmethod
    def _apply_inline_formatting(para) -> str:
        """Apply bold/italic markdown formatting to runs."""
        result = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            result.append(text)
        return ''.join(result) or para.text

    @staticmethod
    def _extract_images(doc: Document, temp_dir: str) -> List[ImageContent]:
        """Extract embedded images from DOCX with size metadata."""
        images = []

        for img_index, rel in enumerate(doc.part.rels.values()):
            if "image" not in rel.target_ref:
                continue

            try:
                blob = rel.target_part.blob
                original_filename = rel.target_ref.split('/')[-1]
                ext = original_filename.split('.')[-1] if '.' in original_filename else 'png'

                # Use index prefix to avoid duplicate filenames
                filename = f"img{img_index}_{original_filename}"
                image_path = os.path.join(temp_dir, filename)

                with open(image_path, 'wb') as f:
                    f.write(blob)

                # Extract image dimensions using PIL
                size = DOCExtractor._get_image_size(blob)

                images.append(ImageContent(path=image_path, format=ext, size=size))
            except (IOError, OSError, ValueError, KeyError) as e:
                logger.warning(f"Failed to extract image {img_index} from DOCX: {e}")

        return images

    @staticmethod
    def _get_image_size(blob: bytes) -> Optional[Tuple[int, int]]:
        """Extract image dimensions from blob data."""
        try:
            with Image.open(io.BytesIO(blob)) as img:
                return img.size
        except (IOError, OSError, ValueError):
            return None

    @staticmethod
    def _extract_tables(doc: Document) -> List[TableContent]:
        """Extract tables from DOCX with markdown conversion."""
        tables = []

        for table_index, table in enumerate(doc.tables):
            try:
                if not table.rows:
                    continue

                # Extract headers, handling merged cells by deduplicating
                headers = DOCExtractor._get_unique_row_cells(table.rows[0])
                if not headers:
                    continue

                # Extract data rows
                rows = []
                for row in table.rows[1:]:
                    cell_values = DOCExtractor._get_unique_row_cells(row)
                    row_data = {}
                    for i, value in enumerate(cell_values):
                        if i < len(headers):
                            row_data[headers[i]] = value
                    rows.append(row_data)

                markdown = DOCExtractor._table_to_markdown(headers, rows)
                tables.append(TableContent(data=rows, markdown=markdown))
            except (ValueError, AttributeError, IndexError) as e:
                logger.warning(f"Failed to extract table {table_index}: {e}")

        return tables

    @staticmethod
    def _get_unique_row_cells(row) -> List[str]:
        """Extract cell values from a row, handling merged cells."""
        seen_cells = set()
        values = []
        for cell in row.cells:
            # Skip duplicate cell references (merged cells)
            cell_id = id(cell._tc)
            if cell_id in seen_cells:
                continue
            seen_cells.add(cell_id)
            values.append(cell.text.strip())
        return values

    @staticmethod
    def _table_to_markdown(headers: List[str], rows: List[Dict[str, str]]) -> str:
        """Convert table data to markdown format."""
        if not headers:
            return ""

        md = "| " + " | ".join(headers) + " |\n"
        md += "| " + " | ".join(["---"] * len(headers)) + " |\n"

        for row in rows:
            values = [str(row.get(h, '')) for h in headers]
            md += "| " + " | ".join(values) + " |\n"

        return md
