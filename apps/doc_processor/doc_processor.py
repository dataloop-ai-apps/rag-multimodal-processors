"""
DOC/DOCX processor app.

DOCX processor with all extraction and processing logic.
"""

import logging
import os
import tempfile
from typing import Dict, Any, List

import dtlpy as dl
from docx import Document

import transforms
from utils.data_types import ExtractedContent, ImageContent, TableContent

logger = logging.getLogger("rag-preprocessor")


class DOCProcessor(dl.BaseServiceRunner):
    """
    DOCX processing application.

    Supports:
    - Text extraction from DOC/DOCX files
    - Table extraction
    - Multiple chunking strategies
    - Text cleaning and normalization
    """

    def __init__(self):
        """Initialize DOC processor."""
        # Configure Dataloop client timeouts
        dl.client_api._upload_session_timeout = 60
        dl.client_api._upload_chuck_timeout = 30

    @staticmethod
    def extract_docx(item: dl.Item, config: Dict[str, Any]) -> ExtractedContent:
        """
        Extract content from DOCX

        Args:
            item: Dataloop DOCX item to extract from
            config: Processing configuration dict

        Returns:
            ExtractedContent: Extracted content with text, images, and tables
        """
        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = item.download(local_path=temp_dir)
            doc = Document(file_path)

            result = ExtractedContent()

            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            result.text = '\n\n'.join(paragraphs)

            # Extract images if requested
            if config.get('extract_images', True):
                result.images = DOCProcessor._extract_images(doc, temp_dir)

            # Extract tables if requested
            if config.get('extract_tables', False):
                result.tables = DOCProcessor._extract_tables(doc)

            result.metadata = {
                'paragraph_count': len(paragraphs),
                'source_file': item.name,
                'image_count': len(result.images),
                'table_count': len(result.tables),
                'processor': 'doc',
            }

            return result

    @staticmethod
    def _extract_images(doc, temp_dir: str) -> List[ImageContent]:
        """Extract embedded images from .docx"""
        images = []

        try:
            # .docx images are in doc.part.rels
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_path = os.path.join(temp_dir, rel.target_ref.split('/')[-1])
                    with open(image_path, 'wb') as f:
                        f.write(rel.target_part.blob)

                    # Extract file extension
                    ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else None

                    images.append(
                        ImageContent(path=image_path, format=ext, caption=None, page_number=None, bbox=None, size=None)
                    )
        except Exception as e:
            logger.warning(f"Failed to extract images from .docx: {e}")

        return images

    @staticmethod
    def _extract_tables(doc) -> List[TableContent]:
        """Extract tables from .docx"""
        tables = []

        for table in doc.tables:
            try:
                # Convert to list of dicts
                rows = []
                headers = [cell.text for cell in table.rows[0].cells]

                for row in table.rows[1:]:
                    row_data = {headers[i]: cell.text for i, cell in enumerate(row.cells)}
                    rows.append(row_data)

                # Convert to markdown
                markdown = DOCProcessor._table_to_markdown(headers, rows)

                tables.append(TableContent(data=rows, markdown=markdown))
            except Exception as e:
                logger.warning(f"Failed to extract table: {e}")

        return tables

    @staticmethod
    def _table_to_markdown(headers: List[str], rows: List[Dict]) -> str:
        """Convert table to markdown format"""
        md = "| " + " | ".join(headers) + " |\n"
        md += "| " + " | ".join(["---"] * len(headers)) + " |\n"

        for row in rows:
            md += "| " + " | ".join([str(row.get(h, '')) for h in headers]) + " |\n"

        return md

    @staticmethod
    def extract(data: Dict[str, Any], config: Dict[str, Any]) -> Dict[str, Any]:
        """Extract content from DOCX file."""
        item = data.get('item')
        if not item:
            raise ValueError("Missing 'item' in data")

        extracted = DOCProcessor.extract_docx(item, config)
        data.update(extracted.to_dict())
        return data

    @staticmethod
    def clean(data: Dict[str, Any], config: Dict[str, Any]) -> Dict[str, Any]:
        """Clean and normalize text."""
        data = transforms.clean_text(data, config)
        data = transforms.normalize_whitespace(data, config)
        return data

    @staticmethod
    def chunk(data: Dict[str, Any], config: Dict[str, Any]) -> Dict[str, Any]:
        """Chunk content based on strategy."""
        strategy = config.get('chunking_strategy', 'recursive')
        link_images = config.get('link_images_to_chunks', True)
        embed_images = config.get('embed_images_in_chunks', False)
        has_images = len(data.get('images', [])) > 0

        if strategy == 'recursive' and has_images:
            if embed_images:
                data = transforms.chunk_with_embedded_images(data, config)
            elif link_images:
                data = transforms.chunk_recursive_with_images(data, config)
            else:
                data = transforms.chunk_text(data, config)
        elif strategy == 'semantic':
            data = transforms.llm_chunk_semantic(data, config)
        else:
            data = transforms.chunk_text(data, config)
        return data

    @staticmethod
    def upload(data: Dict[str, Any], config: Dict[str, Any]) -> Dict[str, Any]:
        """Upload chunks to Dataloop."""
        return transforms.upload_to_dataloop(data, config)

    @staticmethod
    def process_document(item: dl.Item, target_dataset: dl.Dataset, context: dl.Context) -> List[dl.Item]:
        """Dataloop pipeline entry point."""
        config = context.node.metadata.get('customNodeConfig', {})
        return DOCProcessor.run(item, target_dataset, config)

    @staticmethod
    def run(item: dl.Item, target_dataset: dl.Dataset, config: Dict[str, Any]) -> List[dl.Item]:
        """
        Process a DOCX document into chunks.

        Args:
            item: DOCX item to process
            target_dataset: Target dataset for storing chunks
            config: Processing configuration dict

        Returns:
            List of uploaded chunk items
        """
        try:
            data = {'item': item, 'target_dataset': target_dataset}
            data = DOCProcessor.extract(data, config)
            data = DOCProcessor.clean(data, config)
            data = DOCProcessor.chunk(data, config)
            data = DOCProcessor.upload(data, config)

            uploaded = data.get('uploaded_items', [])
            logger.info(f"Processed {item.name}: {len(uploaded)} chunks")
            return uploaded

        except Exception as e:
            logger.error(f"Processing failed: {str(e)}", exc_info=True)
            raise
